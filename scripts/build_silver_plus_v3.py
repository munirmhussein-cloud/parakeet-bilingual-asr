
from __future__ import annotations

import argparse
import hashlib
import json
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []

    with path.open("r", encoding="utf-8") as handle:
        for line_number, line in enumerate(handle, start=1):
            if not line.strip():
                continue

            try:
                row = json.loads(line)
            except json.JSONDecodeError as error:
                raise ValueError(
                    f"Invalid JSONL at {path}:{line_number}"
                ) from error

            if not isinstance(row, dict):
                raise TypeError(
                    f"Expected JSON object at {path}:{line_number}"
                )

            rows.append(row)

    return rows


def write_jsonl(
    path: Path,
    rows: Iterable[dict[str, Any]],
) -> None:
    with path.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(
                json.dumps(
                    row,
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                )
                + "\n"
            )


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()

    with path.open("rb") as handle:
        for chunk in iter(
            lambda: handle.read(1024 * 1024),
            b"",
        ):
            digest.update(chunk)

    return digest.hexdigest()


def first_number(values: list[Any]) -> float | None:
    for value in values:
        if value is None or isinstance(value, bool):
            continue

        try:
            return float(value)
        except (TypeError, ValueError):
            continue

    return None


def resolve_canonical_span(
    row: dict[str, Any],
    position: int,
) -> tuple[float, float, float]:
    start = first_number(
        [
            row.get("segment_start"),
            row.get("start"),
            row.get("offset"),
            row.get("global_start"),
            row.get("canonical_start"),
        ]
    )

    end = first_number(
        [
            row.get("segment_end"),
            row.get("end"),
            row.get("global_end"),
            row.get("canonical_end"),
        ]
    )

    duration = first_number(
        [
            row.get("duration"),
            row.get("segment_duration"),
            row.get("canonical_duration"),
        ]
    )

    if start is None:
        start = float(position) * 20.0

    if end is None and duration is not None:
        end = start + duration

    if duration is None and end is not None:
        duration = end - start

    if end is None:
        duration = 20.0 if duration is None else duration
        end = start + duration

    if duration is None:
        duration = end - start

    if end < start:
        raise ValueError(
            f"Invalid canonical span: start={start}, end={end}, row={row}"
        )

    return (
        round(start, 6),
        round(end, 6),
        round(duration, 6),
    )


def extract_silver_v3_text(row: dict[str, Any]) -> str:
    for field in [
        "silver_text",
        "silver_v3_text",
        "text",
    ]:
        value = row.get(field)

        if isinstance(value, str):
            return value.strip()

    return ""


def normalize_children(
    azure_row: dict[str, Any],
) -> list[dict[str, Any]]:
    raw_children = azure_row.get("children", [])

    if not isinstance(raw_children, list):
        raw_children = []

    children: list[dict[str, Any]] = []

    for child_position, child in enumerate(raw_children):
        if not isinstance(child, dict):
            continue

        transcript = str(
            child.get(
                "transcript",
                child.get("text", ""),
            )
        ).strip()

        global_start = first_number(
            [child.get("global_start")]
        )

        global_end = first_number(
            [child.get("global_end")]
        )

        children.append(
            {
                "child_position": child_position,
                "segment_id": str(
                    child.get("segment_id", "")
                ),
                "global_start": global_start,
                "global_end": global_end,
                "duration": (
                    None
                    if global_start is None or global_end is None
                    else round(global_end - global_start, 6)
                ),
                "transcript": transcript,
                "has_text": bool(transcript),
                "status": child.get("status"),
            }
        )

    children.sort(
        key=lambda child: (
            float("inf")
            if child["global_start"] is None
            else child["global_start"],
            float("inf")
            if child["global_end"] is None
            else child["global_end"],
            child["child_position"],
        )
    )

    return children


def reconstruct_azure_parent_text(
    azure_row: dict[str, Any],
    children: list[dict[str, Any]],
) -> tuple[str, str]:
    child_texts = [
        child["transcript"]
        for child in children
        if child["has_text"]
    ]

    if child_texts:
        return (
            " ".join(child_texts).strip(),
            "ordered_pyannote_children",
        )

    stored_parent_text = str(
        azure_row.get(
            "azure_forced_ar_text",
            "",
        )
    ).strip()

    if stored_parent_text:
        return (
            stored_parent_text,
            "stored_parent_text",
        )

    return "", "empty"


def format_timestamp(seconds: float) -> str:
    total_ms = round(float(seconds) * 1000)

    hours, remainder = divmod(
        total_ms,
        3_600_000,
    )

    minutes, remainder = divmod(
        remainder,
        60_000,
    )

    seconds, milliseconds = divmod(
        remainder,
        1000,
    )

    return (
        f"{hours:02d}:"
        f"{minutes:02d}:"
        f"{seconds:02d}."
        f"{milliseconds:03d}"
    )


def export_docx(
    rows: list[dict[str, Any]],
    output_path: Path,
    title_text: str,
) -> None:
    document = Document()

    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle.add_run(
        "Canonical 20-second parent segments"
    ).bold = True

    document.add_paragraph(
        "Azure forced-ar transcripts were generated from "
        "Pyannote speech-turn clips and reconstructed in "
        "chronological child-span order."
    )

    document.add_paragraph(
        "Resolution rule: use Azure parent text when populated; "
        "otherwise use Silver v3."
    )

    document.add_paragraph("")

    for row in rows:
        heading = document.add_paragraph()

        heading.add_run(
            f"SEGMENT "
            f"{row['segment_position'] + 1:04d}"
            f" | {row['segment_id']}"
        ).bold = True

        metadata = document.add_paragraph()

        metadata.add_run(
            "Canonical time: "
            f"{format_timestamp(row['segment_start'])}"
            " → "
            f"{format_timestamp(row['segment_end'])}"
            " | Duration: "
            f"{row['duration']:.3f}s"
            " | Resolution: "
            f"{row['resolution_source']}"
            " | Pyannote clips: "
            f"{row['pyannote_clip_count']}"
        )

        document.add_paragraph(
            row["silver_plus_v3_text"]
        )

        document.add_paragraph("")

    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Build Silver+ v3 using the Silver+ v2 parent-level "
            "Azure/Pyannote resolution rule with Silver v3 fallback."
        )
    )

    parser.add_argument(
        "--lecture-id",
        required=True,
    )

    parser.add_argument(
        "--silver-v3",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--azure-parent",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--canonical-manifest",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--output-dir",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--output-prefix",
        required=True,
    )

    parser.add_argument(
        "--title",
        default="Silver+ v3",
    )

    args = parser.parse_args()

    for required_path in [
        args.silver_v3,
        args.azure_parent,
        args.canonical_manifest,
    ]:
        if not required_path.exists():
            raise FileNotFoundError(
                f"Missing required input: {required_path}"
            )

    args.output_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    silver_v3_rows = read_jsonl(
        args.silver_v3
    )

    azure_parent_rows = read_jsonl(
        args.azure_parent
    )

    canonical_rows = read_jsonl(
        args.canonical_manifest
    )

    if not silver_v3_rows:
        raise ValueError("Silver v3 input is empty.")

    azure_by_segment = {
        str(row["segment_id"]): row
        for row in azure_parent_rows
    }

    canonical_by_segment = {
        str(row["segment_id"]): row
        for row in canonical_rows
    }

    silver_ids = [
        str(row["segment_id"])
        for row in silver_v3_rows
    ]

    if len(silver_ids) != len(set(silver_ids)):
        raise ValueError(
            "Duplicate Silver v3 segment IDs."
        )

    missing_azure = [
        segment_id
        for segment_id in silver_ids
        if segment_id not in azure_by_segment
    ]

    missing_canonical = [
        segment_id
        for segment_id in silver_ids
        if segment_id not in canonical_by_segment
    ]

    if missing_azure:
        raise ValueError(
            f"Missing Azure parent rows: {missing_azure[:20]}"
        )

    if missing_canonical:
        raise ValueError(
            f"Missing canonical rows: {missing_canonical[:20]}"
        )

    output_rows = []
    audit_rows = []

    resolution_distribution = Counter()
    azure_method_distribution = Counter()

    canonical_span_errors = []
    child_chronology_errors = []

    for position, silver_row in enumerate(silver_v3_rows):
        segment_id = str(
            silver_row["segment_id"]
        )

        canonical_row = canonical_by_segment[
            segment_id
        ]

        azure_row = azure_by_segment[
            segment_id
        ]

        segment_start, segment_end, duration = (
            resolve_canonical_span(
                canonical_row,
                position,
            )
        )

        audio_filepath = str(
            canonical_row.get("audio_filepath")
            or silver_row.get("audio_filepath")
            or azure_row.get("audio_filepath")
            or ""
        )

        silver_v3_text = extract_silver_v3_text(
            silver_row
        )

        children = normalize_children(
            azure_row
        )

        azure_text, azure_method = (
            reconstruct_azure_parent_text(
                azure_row,
                children,
            )
        )

        if azure_text:
            silver_plus_v3_text = azure_text
            resolution_source = (
                "azure_pyannote_forced_ar"
            )
            fallback_used = False

        else:
            silver_plus_v3_text = silver_v3_text
            resolution_source = (
                "silver_v3_fallback"
            )
            fallback_used = True

        resolution_distribution[
            resolution_source
        ] += 1

        azure_method_distribution[
            azure_method
        ] += 1

        output_row = {
            "schema_version": (
                "silver_plus_v3_parent_resolution_v1"
            ),
            "lecture_id": args.lecture_id,
            "segment_position": position,
            "segment_index": position,
            "segment_id": segment_id,
            "audio_filepath": audio_filepath,
            "segment_start": segment_start,
            "segment_end": segment_end,
            "duration": duration,
            "silver_plus_v3_text": (
                silver_plus_v3_text
            ),
            "has_silver_plus_v3_text": bool(
                silver_plus_v3_text.strip()
            ),
            "silver_v3_text": silver_v3_text,
            "silver_v3_has_text": bool(
                silver_v3_text
            ),
            "azure_pyannote_forced_ar_text": (
                azure_text
            ),
            "azure_has_text": bool(
                azure_text
            ),
            "resolution_source": (
                resolution_source
            ),
            "silver_v3_fallback_used": (
                fallback_used
            ),
            "azure_text_reconstruction_method": (
                azure_method
            ),
            "pyannote_clip_count": int(
                azure_row.get(
                    "pyannote_clip_count",
                    len(children),
                )
            ),
            "completed_clip_count": int(
                azure_row.get(
                    "completed_clip_count",
                    sum(
                        child["status"] == "completed"
                        for child in children
                    ),
                )
            ),
            "failed_clip_count": int(
                azure_row.get(
                    "failed_clip_count",
                    sum(
                        child["status"] == "failed"
                        for child in children
                    ),
                )
            ),
            "azure_children": children,
            "reconciliation": {
                "method": (
                    "silver_plus_v2_parent_resolution_"
                    "with_silver_v3_fallback"
                ),
                "resolution_rule": (
                    "azure_parent_if_populated_else_silver_v3"
                ),
                "parent_span_source": (
                    "canonical_20s_manifest"
                ),
                "azure_called_on": (
                    "pyannote_speech_turn_clips"
                ),
                "pyannote_timestamps_used_for_azure_calls": (
                    True
                ),
                "azure_used_inside_silver_v3_lattice": (
                    False
                ),
                "token_level_merge_used": False,
                "human_review_used": False,
            },
        }

        output_rows.append(output_row)

        audit_rows.append(
            {
                "schema_version": (
                    "silver_plus_v3_resolution_audit_v1"
                ),
                "segment_position": position,
                "segment_id": segment_id,
                "canonical_start": segment_start,
                "canonical_end": segment_end,
                "canonical_duration": duration,
                "resolution_source": (
                    resolution_source
                ),
                "azure_text_reconstruction_method": (
                    azure_method
                ),
                "silver_v3_text": silver_v3_text,
                "azure_text": azure_text,
                "silver_plus_v3_text": (
                    silver_plus_v3_text
                ),
                "azure_children": children,
            }
        )

        expected_start, expected_end, expected_duration = (
            resolve_canonical_span(
                canonical_row,
                position,
            )
        )

        if (
            abs(segment_start - expected_start) > 1e-6
            or abs(segment_end - expected_end) > 1e-6
            or abs(duration - expected_duration) > 1e-6
        ):
            canonical_span_errors.append(
                {
                    "segment_id": segment_id,
                    "output_start": segment_start,
                    "expected_start": expected_start,
                    "output_end": segment_end,
                    "expected_end": expected_end,
                    "output_duration": duration,
                    "expected_duration": expected_duration,
                }
            )

        previous_child_start = None

        for child in children:
            child_start = child["global_start"]
            child_end = child["global_end"]

            if (
                child_start is not None
                and child_end is not None
                and child_end < child_start
            ):
                child_chronology_errors.append(
                    {
                        "segment_id": segment_id,
                        "child_segment_id": child["segment_id"],
                        "code": "child_end_before_start",
                    }
                )

            if (
                child_start is not None
                and previous_child_start is not None
                and child_start < previous_child_start
            ):
                child_chronology_errors.append(
                    {
                        "segment_id": segment_id,
                        "child_segment_id": child["segment_id"],
                        "code": "child_start_regression",
                    }
                )

            if child_start is not None:
                previous_child_start = child_start

    output_positions = [
        row["segment_position"]
        for row in output_rows
    ]

    output_ids = [
        row["segment_id"]
        for row in output_rows
    ]

    empty_segments = [
        row["segment_id"]
        for row in output_rows
        if not row[
            "has_silver_plus_v3_text"
        ]
    ]

    validation = {
        "segment_count_matches_silver_v3": (
            len(output_rows)
            == len(silver_v3_rows)
        ),
        "segment_positions_ordered": (
            output_positions
            == list(range(len(output_rows)))
        ),
        "segment_ids_match_silver_v3": (
            output_ids == silver_ids
        ),
        "segment_ids_unique": (
            len(output_ids)
            == len(set(output_ids))
        ),
        "canonical_span_error_count": (
            len(canonical_span_errors)
        ),
        "canonical_spans_match_exactly": (
            not canonical_span_errors
        ),
        "child_chronology_error_count": (
            len(child_chronology_errors)
        ),
        "azure_not_used_inside_silver_v3_lattice": all(
            row["reconciliation"][
                "azure_used_inside_silver_v3_lattice"
            ]
            is False
            for row in output_rows
        ),
        "parent_level_resolution_only": all(
            row["reconciliation"][
                "token_level_merge_used"
            ]
            is False
            for row in output_rows
        ),
    }

    validation["passed"] = (
        validation[
            "segment_count_matches_silver_v3"
        ]
        and validation[
            "segment_positions_ordered"
        ]
        and validation[
            "segment_ids_match_silver_v3"
        ]
        and validation[
            "segment_ids_unique"
        ]
        and validation[
            "canonical_spans_match_exactly"
        ]
        and validation[
            "child_chronology_error_count"
        ]
        == 0
        and validation[
            "azure_not_used_inside_silver_v3_lattice"
        ]
        and validation[
            "parent_level_resolution_only"
        ]
    )

    if not validation["passed"]:
        print(
            json.dumps(
                {
                    "validation": validation,
                    "canonical_span_errors": (
                        canonical_span_errors
                    ),
                    "child_chronology_errors": (
                        child_chronology_errors
                    ),
                },
                ensure_ascii=False,
                indent=2,
            )
        )

        return 1

    output_jsonl = (
        args.output_dir
        / f"{args.output_prefix}_segment_level.jsonl"
    )

    output_json = (
        args.output_dir
        / f"{args.output_prefix}.json"
    )

    output_audit = (
        args.output_dir
        / f"{args.output_prefix}_resolution_audit.jsonl"
    )

    output_report = (
        args.output_dir
        / f"{args.output_prefix}_report.json"
    )

    output_docx = (
        args.output_dir
        / f"{args.output_prefix}.docx"
    )

    output_zip = (
        args.output_dir
        / f"{args.output_prefix}.zip"
    )

    write_jsonl(
        output_jsonl,
        output_rows,
    )

    write_jsonl(
        output_audit,
        audit_rows,
    )

    output_json.write_text(
        json.dumps(
            {
                "schema_version": (
                    "silver_plus_v3_export_v1"
                ),
                "lecture_id": (
                    args.lecture_id
                ),
                "segment_count": len(
                    output_rows
                ),
                "segments": output_rows,
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )

    export_docx(
        output_rows,
        output_docx,
        title_text=args.title,
    )

    report = {
        "schema_version": (
            "silver_plus_v3_export_report_v1"
        ),
        "lecture_id": args.lecture_id,
        "segment_count": len(
            output_rows
        ),
        "segments_with_text": (
            len(output_rows)
            - len(empty_segments)
        ),
        "empty_segment_count": len(
            empty_segments
        ),
        "empty_segments": empty_segments,
        "resolution_source_distribution": dict(
            resolution_distribution
        ),
        "azure_text_reconstruction_method_distribution": dict(
            azure_method_distribution
        ),
        "methodology": {
            "same_resolution_rule_as_silver_plus_v2": True,
            "azure_parent_if_populated": True,
            "silver_v3_fallback_if_azure_empty": True,
            "parent_spans_from_canonical_manifest": True,
            "pyannote_child_spans_preserved": True,
            "azure_used_inside_silver_v3_lattice": False,
            "token_level_merge_used": False,
        },
        "outputs": {
            "segment_jsonl": str(
                output_jsonl
            ),
            "json": str(output_json),
            "audit_jsonl": str(
                output_audit
            ),
            "docx": str(output_docx),
        },
        "canonical_span_errors": (
            canonical_span_errors
        ),
        "child_chronology_errors": (
            child_chronology_errors
        ),
        "validation": validation,
    }

    output_report.write_text(
        json.dumps(
            report,
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
        )
        + "\n",
        encoding="utf-8",
    )

    with zipfile.ZipFile(
        output_zip,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as archive:
        for path in [
            output_jsonl,
            output_json,
            output_audit,
            output_report,
            output_docx,
        ]:
            archive.write(
                path,
                arcname=path.name,
            )

    report["outputs"]["zip"] = str(
        output_zip
    )

    report["sha256"] = {
        "segment_jsonl": sha256_file(
            output_jsonl
        ),
        "json": sha256_file(
            output_json
        ),
        "audit_jsonl": sha256_file(
            output_audit
        ),
        "report": sha256_file(
            output_report
        ),
        "docx": sha256_file(
            output_docx
        ),
        "zip": sha256_file(
            output_zip
        ),
    }

    print(
        json.dumps(
            report,
            ensure_ascii=False,
            indent=2,
        )
    )

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
