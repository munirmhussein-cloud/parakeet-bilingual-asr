#!/usr/bin/env python3
"""Run Faster-Whisper Turbo on a whole lecture and project words into canonical 20s rows.

Bronze v2.4 keeps one whole-lecture Whisper decode for speed and context. After
inference, word timestamps are assigned exactly once to deterministic contiguous
20-second windows. This removes repeated display timestamps while producing a
Silver-v3-like canonical segment-level JSONL shape.
"""

from __future__ import annotations

import argparse
import json
import math
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

SCHEMA_VERSION = "bronze_v2_4_canonical_20s_v1"
PIPELINE_VERSION = "bronze_v2.4"
WINDOW_SECONDS = 20.0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--lecture-id", required=True)
    parser.add_argument("--audio", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--model", default="turbo")
    parser.add_argument("--compute-type", default="float16")
    parser.add_argument("--beam-size", type=int, default=5)
    parser.add_argument("--best-of", type=int, default=5)
    parser.add_argument("--temperature", type=float, default=0.0)
    parser.add_argument("--download-root", type=Path, default=None)
    parser.add_argument("--overwrite", action="store_true")
    return parser.parse_args()


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def normalize_word(value: Any) -> str:
    return " ".join(str(value or "").strip().split())


def join_words(words: list[str]) -> str:
    text = "".join(words).strip()
    return " ".join(text.split())


def timestamp(seconds: float) -> str:
    total = max(0, int(seconds))
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def load_whole_words(payload: dict[str, Any]) -> list[dict[str, Any]]:
    words: list[dict[str, Any]] = []
    ordinal = 0
    for segment in payload.get("segments", []):
        segment_words = segment.get("words") or []
        if segment_words:
            for word in segment_words:
                text = str(word.get("word") or "")
                if not text.strip():
                    continue
                start = float(word.get("start", segment.get("start", 0.0)))
                end = float(word.get("end", start))
                words.append({
                    "ordinal": ordinal,
                    "start": start,
                    "end": max(start, end),
                    "word": text,
                    "probability": word.get("probability"),
                })
                ordinal += 1
        else:
            text = normalize_word(segment.get("text"))
            if text:
                start = float(segment.get("start", 0.0))
                end = float(segment.get("end", start))
                words.append({
                    "ordinal": ordinal,
                    "start": start,
                    "end": max(start, end),
                    "word": " " + text,
                    "probability": None,
                })
                ordinal += 1
    return words


def project_canonical_rows(
    lecture_id: str,
    audio: Path,
    duration: float,
    words: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    count = max(1, math.ceil(duration / WINDOW_SECONDS))
    buckets: list[list[dict[str, Any]]] = [[] for _ in range(count)]

    for word in words:
        midpoint = (float(word["start"]) + float(word["end"])) / 2.0
        index = min(count - 1, max(0, int(midpoint // WINDOW_SECONDS)))
        buckets[index].append(word)

    rows: list[dict[str, Any]] = []
    assigned_ordinals: list[int] = []
    for index, bucket in enumerate(buckets):
        start = index * WINDOW_SECONDS
        end = min(duration, (index + 1) * WINDOW_SECONDS)
        text = join_words([item["word"] for item in bucket])
        assigned_ordinals.extend(item["ordinal"] for item in bucket)
        rows.append({
            "schema_version": SCHEMA_VERSION,
            "pipeline_version": PIPELINE_VERSION,
            "lecture_id": lecture_id,
            "audio_filepath": str(audio.resolve()),
            "segment_id": f"{lecture_id}__canonical_20s__{index:05d}",
            "segment_index": index,
            "segment_position": index,
            "segment_start": start,
            "segment_end": end,
            "duration": max(0.0, end - start),
            "has_bronze_text": bool(text),
            "bronze_text": text,
            "token_count": len(text.split()),
            "word_count": len(bucket),
            "word_start_ordinal": bucket[0]["ordinal"] if bucket else None,
            "word_end_ordinal": bucket[-1]["ordinal"] if bucket else None,
            "words": bucket,
        })

    expected = list(range(len(words)))
    if assigned_ordinals != expected:
        raise RuntimeError("Canonical projection did not assign every word exactly once")
    return rows


def write_docx(path: Path, lecture_id: str, rows: list[dict[str, Any]], metadata: dict[str, Any]) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX export") from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    title = doc.add_heading(lecture_id.replace("_", " ").title(), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Bronze v2.4 — Faster-Whisper Turbo — Canonical 20s").bold = True
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(
        f"{len(rows):,} canonical rows | {metadata['word_count']:,} words | "
        f"RTF {metadata['runtime'].get('real_time_factor', 0):.4f}"
    ).italic = True
    doc.add_paragraph()

    for row in rows:
        paragraph = doc.add_paragraph()
        label = f"[{timestamp(row['segment_start'])} – {timestamp(row['segment_end'])}] "
        run = paragraph.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        paragraph.add_run(row["bronze_text"] or "[No speech]")

    doc.core_properties.title = f"{lecture_id} Bronze v2.4 Transcript"
    doc.core_properties.subject = "Faster-Whisper Turbo canonical 20-second transcript"
    doc.core_properties.author = "Parakeet Bilingual ASR Project"
    doc.save(path)


def main() -> int:
    args = parse_args()
    audio = args.audio.expanduser().resolve()
    output_dir = args.output_dir.expanduser().resolve()
    if not audio.is_file():
        raise FileNotFoundError(audio)
    output_dir.mkdir(parents=True, exist_ok=True)

    whole_json = output_dir / f"{args.lecture_id}_bronze_v2_4_faster_whisper_turbo_whole.json"
    canonical_json = output_dir / f"{args.lecture_id}_bronze_v2_4_canonical_20s.json"
    canonical_jsonl = output_dir / f"{args.lecture_id}_bronze_v2_4_canonical_20s.jsonl"
    docx_path = output_dir / f"{args.lecture_id}_bronze_v2_4_canonical_20s.docx"
    manifest_path = output_dir / "PACKAGE_MANIFEST.json"

    existing = [whole_json, canonical_json, canonical_jsonl, docx_path, manifest_path]
    if not args.overwrite and any(path.exists() for path in existing):
        raise FileExistsError("Output package exists; pass --overwrite")

    v23_runner = Path(__file__).with_name("run_bronze_v2_3_faster_whisper.py")
    command = [
        sys.executable, str(v23_runner),
        "--lecture-id", args.lecture_id,
        "--audio", str(audio),
        "--output", str(whole_json),
        "--model", args.model,
        "--device", "cuda",
        "--compute-type", args.compute_type,
        "--beam-size", str(args.beam_size),
        "--best-of", str(args.best_of),
        "--temperature", str(args.temperature),
        "--num-workers", "1",
        "--overwrite",
    ]
    if args.download_root:
        command.extend(["--download-root", str(args.download_root.expanduser().resolve())])
    subprocess.run(command, check=True)

    whole = json.loads(whole_json.read_text(encoding="utf-8"))
    duration = float(whole.get("duration") or 0.0)
    words = load_whole_words(whole)
    if not words:
        raise RuntimeError("Whole-lecture output contains no timestamped words")
    rows = project_canonical_rows(args.lecture_id, audio, duration, words)

    package = {
        "schema_version": SCHEMA_VERSION,
        "pipeline_version": PIPELINE_VERSION,
        "lecture_id": args.lecture_id,
        "created_at": utc_now(),
        "source_audio": whole.get("source_audio"),
        "engine": whole.get("engine"),
        "runtime": whole.get("runtime"),
        "duration": duration,
        "window_seconds": WINDOW_SECONDS,
        "segment_count": len(rows),
        "nonempty_segment_count": sum(1 for row in rows if row["bronze_text"]),
        "word_count": len(words),
        "text": " ".join(row["bronze_text"] for row in rows if row["bronze_text"]),
        "segments": rows,
    }
    canonical_json.write_text(json.dumps(package, ensure_ascii=False, indent=2), encoding="utf-8")
    with canonical_jsonl.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")
    write_docx(docx_path, args.lecture_id, rows, package)

    manifest = {
        "schema_version": "bronze_v2_4_package_manifest_v1",
        "pipeline_version": PIPELINE_VERSION,
        "lecture_id": args.lecture_id,
        "created_at": utc_now(),
        "files": {
            "whole_json": whole_json.name,
            "canonical_json": canonical_json.name,
            "canonical_jsonl": canonical_jsonl.name,
            "docx": docx_path.name,
        },
        "metrics": {
            "duration": duration,
            "canonical_segment_count": len(rows),
            "nonempty_segment_count": package["nonempty_segment_count"],
            "word_count": len(words),
            "elapsed_seconds": whole.get("runtime", {}).get("elapsed_seconds"),
            "real_time_factor": whole.get("runtime", {}).get("real_time_factor"),
        },
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print("\n=== BRONZE V2.4 COMPLETE ===")
    print(f"Output directory: {output_dir}")
    print(f"Canonical rows: {len(rows):,}")
    print(f"Non-empty rows: {package['nonempty_segment_count']:,}")
    print(f"Words assigned exactly once: {len(words):,}")
    print(f"JSON: {canonical_json}")
    print(f"JSONL: {canonical_jsonl}")
    print(f"DOCX: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
