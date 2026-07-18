#!/usr/bin/env python3
"""Export a completed Faster-Whisper whole JSON into Bronze v2.4 canonical 20s files.

The exporter prefers word timestamps. If the serialized Faster-Whisper output has
no word rows, it falls back to Whisper segment timestamps. Each source unit is
assigned exactly once by temporal midpoint, preventing duplicated text and
repeated source timestamps. No ASR inference is performed by this script.
"""

from __future__ import annotations

import argparse
import json
import math
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

WINDOW_SECONDS = 20.0
SCHEMA_VERSION = "bronze_v2_4_canonical_20s_v2"
PIPELINE_VERSION = "bronze_v2.4"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--lecture-id", required=True)
    parser.add_argument("--audio", required=True, type=Path)
    parser.add_argument("--whole-json", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--overwrite", action="store_true")
    return parser.parse_args()


def clean_text(value: Any) -> str:
    return " ".join(str(value or "").replace("\u00a0", " ").split())


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def ts(seconds: float) -> str:
    total = max(0, int(seconds))
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def extract_source_units(payload: dict[str, Any]) -> tuple[list[dict[str, Any]], str]:
    segments = payload.get("segments")
    if not isinstance(segments, list):
        raise RuntimeError("Whole JSON does not contain a segment list")

    word_units: list[dict[str, Any]] = []
    ordinal = 0
    for segment_index, segment in enumerate(segments):
        if not isinstance(segment, dict):
            continue
        words = segment.get("words")
        if not isinstance(words, list):
            continue
        for word in words:
            if not isinstance(word, dict):
                continue
            text = str(word.get("word") or "")
            if not text.strip():
                continue
            start = float(word.get("start", segment.get("start", 0.0)) or 0.0)
            end = float(word.get("end", start) or start)
            word_units.append({
                "ordinal": ordinal,
                "source_type": "word",
                "source_segment_index": segment_index,
                "start": start,
                "end": max(start, end),
                "text": text,
                "probability": word.get("probability"),
            })
            ordinal += 1

    if word_units:
        return word_units, "word"

    segment_units: list[dict[str, Any]] = []
    for segment_index, segment in enumerate(segments):
        if not isinstance(segment, dict):
            continue
        text = clean_text(segment.get("text"))
        if not text:
            continue
        start = float(segment.get("start", 0.0) or 0.0)
        end = float(segment.get("end", start) or start)
        segment_units.append({
            "ordinal": len(segment_units),
            "source_type": "whisper_segment",
            "source_segment_index": segment_index,
            "start": start,
            "end": max(start, end),
            "text": text,
            "probability": None,
        })

    if segment_units:
        return segment_units, "whisper_segment"

    # Last-resort diagnostic fallback for legacy whole JSONs.
    transcript = clean_text(payload.get("text"))
    duration = float(payload.get("duration") or 0.0)
    if transcript and duration > 0:
        return [{
            "ordinal": 0,
            "source_type": "whole_transcript",
            "source_segment_index": None,
            "start": 0.0,
            "end": duration,
            "text": transcript,
            "probability": None,
        }], "whole_transcript"

    raise RuntimeError(
        "Whole JSON contains no usable word timestamps, segment text, or top-level transcript"
    )


def join_unit_text(units: list[dict[str, Any]], source_type: str) -> str:
    values = [str(unit["text"]) for unit in units]
    if source_type == "word":
        return " ".join("".join(values).split())
    return " ".join(clean_text(value) for value in values if clean_text(value))


def canonicalize(
    lecture_id: str,
    audio: Path,
    duration: float,
    units: list[dict[str, Any]],
    source_type: str,
) -> list[dict[str, Any]]:
    count = max(1, math.ceil(duration / WINDOW_SECONDS))
    buckets: list[list[dict[str, Any]]] = [[] for _ in range(count)]

    for unit in units:
        midpoint = (float(unit["start"]) + float(unit["end"])) / 2.0
        index = max(0, min(count - 1, int(midpoint // WINDOW_SECONDS)))
        buckets[index].append(unit)

    assigned = [unit["ordinal"] for bucket in buckets for unit in bucket]
    expected = list(range(len(units)))
    if assigned != expected:
        raise RuntimeError("Source units were not assigned exactly once in source order")

    rows: list[dict[str, Any]] = []
    for index, bucket in enumerate(buckets):
        start = index * WINDOW_SECONDS
        end = min(duration, (index + 1) * WINDOW_SECONDS)
        text = join_unit_text(bucket, source_type)
        rows.append({
            "schema_version": SCHEMA_VERSION,
            "pipeline_version": PIPELINE_VERSION,
            "lecture_id": lecture_id,
            "audio_filepath": str(audio.resolve()),
            "segment_id": f"{lecture_id}__canonical_20s__{index:05d}",
            "segment_index": index,
            "segment_position": index,
            "segment_start": start,
            "segment_end": end,
            "duration": max(0.0, end - start),
            "has_bronze_text": bool(text),
            "bronze_text": text,
            "token_count": len(text.split()),
            "source_unit_type": source_type,
            "source_unit_count": len(bucket),
            "source_start_ordinal": bucket[0]["ordinal"] if bucket else None,
            "source_end_ordinal": bucket[-1]["ordinal"] if bucket else None,
            "source_units": bucket,
        })
    return rows


def write_docx(path: Path, lecture_id: str, rows: list[dict[str, Any]], package: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    title = document.add_heading(lecture_id.replace("_", " ").title(), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Bronze v2.4 — Faster-Whisper Turbo — Canonical 20s").bold = True
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(
        f"{len(rows):,} canonical rows | {package['word_count']:,} transcript words | "
        f"projection source: {package['projection_source']}"
    ).italic = True
    document.add_paragraph()

    for row in rows:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"[{ts(row['segment_start'])} – {ts(row['segment_end'])}] "
        )
        run.bold = True
        run.font.size = Pt(9)
        paragraph.add_run(row["bronze_text"] or "[No speech]")

    document.core_properties.title = f"{lecture_id} Bronze v2.4 Transcript"
    document.core_properties.subject = "Canonical 20-second Faster-Whisper Turbo transcript"
    document.core_properties.author = "Parakeet Bilingual ASR Project"
    document.save(path)


def main() -> int:
    args = parse_args()
    audio = args.audio.expanduser().resolve()
    whole_json = args.whole_json.expanduser().resolve()
    output_dir = args.output_dir.expanduser().resolve()

    if not audio.is_file():
        raise FileNotFoundError(audio)
    if not whole_json.is_file():
        raise FileNotFoundError(whole_json)
    output_dir.mkdir(parents=True, exist_ok=True)

    canonical_json = output_dir / f"{args.lecture_id}_bronze_v2_4_canonical_20s.json"
    canonical_jsonl = output_dir / f"{args.lecture_id}_bronze_v2_4_canonical_20s.jsonl"
    docx_path = output_dir / f"{args.lecture_id}_bronze_v2_4_canonical_20s.docx"
    manifest_path = output_dir / "PACKAGE_MANIFEST.json"

    outputs = [canonical_json, canonical_jsonl, docx_path, manifest_path]
    if not args.overwrite and any(path.exists() for path in outputs):
        raise FileExistsError("Canonical outputs exist; pass --overwrite")

    whole = json.loads(whole_json.read_text(encoding="utf-8"))
    duration = float(whole.get("duration") or 0.0)
    if duration <= 0:
        raise RuntimeError("Whole JSON has an invalid duration")

    units, source_type = extract_source_units(whole)
    rows = canonicalize(args.lecture_id, audio, duration, units, source_type)
    transcript = " ".join(row["bronze_text"] for row in rows if row["bronze_text"])

    package = {
        "schema_version": SCHEMA_VERSION,
        "pipeline_version": PIPELINE_VERSION,
        "lecture_id": args.lecture_id,
        "created_at": utc_now(),
        "source_audio": whole.get("source_audio"),
        "source_whole_json": str(whole_json),
        "engine": whole.get("engine"),
        "runtime": whole.get("runtime"),
        "duration": duration,
        "window_seconds": WINDOW_SECONDS,
        "projection_source": source_type,
        "source_unit_count": len(units),
        "segment_count": len(rows),
        "nonempty_segment_count": sum(bool(row["bronze_text"]) for row in rows),
        "word_count": len(transcript.split()),
        "text": transcript,
        "segments": rows,
    }

    canonical_json.write_text(json.dumps(package, ensure_ascii=False, indent=2), encoding="utf-8")
    with canonical_jsonl.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")
    write_docx(docx_path, args.lecture_id, rows, package)

    manifest = {
        "schema_version": "bronze_v2_4_package_manifest_v2",
        "pipeline_version": PIPELINE_VERSION,
        "lecture_id": args.lecture_id,
        "created_at": utc_now(),
        "projection_source": source_type,
        "files": {
            "whole_json": whole_json.name,
            "canonical_json": canonical_json.name,
            "canonical_jsonl": canonical_jsonl.name,
            "docx": docx_path.name,
        },
        "metrics": {
            "duration": duration,
            "canonical_segment_count": len(rows),
            "nonempty_segment_count": package["nonempty_segment_count"],
            "source_unit_count": len(units),
            "word_count": package["word_count"],
            "elapsed_seconds": whole.get("runtime", {}).get("elapsed_seconds"),
            "real_time_factor": whole.get("runtime", {}).get("real_time_factor"),
        },
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print("\n=== BRONZE V2.4 CANONICAL EXPORT COMPLETE ===")
    print(f"Projection source: {source_type}")
    print(f"Source units assigned once: {len(units):,}")
    print(f"Canonical rows: {len(rows):,}")
    print(f"Non-empty rows: {package['nonempty_segment_count']:,}")
    print(f"Transcript words: {package['word_count']:,}")
    print(f"JSON: {canonical_json}")
    print(f"JSONL: {canonical_jsonl}")
    print(f"DOCX: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
