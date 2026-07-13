#!/usr/bin/env python3

from __future__ import annotations

import argparse
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SCHEMA_VERSION = "gold_review_package_v1"


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []

    with path.open("r", encoding="utf-8") as handle:
        for line_number, line in enumerate(handle, start=1):
            line = line.strip()

            if not line:
                continue

            try:
                rows.append(json.loads(line))
            except json.JSONDecodeError as exc:
                raise ValueError(
                    f"Invalid JSONL at line {line_number}: {path}"
                ) from exc

    return rows


def safe_name(value: str) -> str:
    return (
        value.replace("/", "_")
        .replace("\\", "_")
        .replace(" ", "_")
        .replace("–", "-")
    )


def format_timestamp(seconds: float | int | None) -> str:
    value = float(seconds or 0.0)
    milliseconds = int(round(value * 1000))

    hours, remainder = divmod(milliseconds, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1_000)

    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"


def resolve_segment_bounds(row: dict[str, Any]) -> tuple[float, float]:
    start = row.get("start")

    if start is None:
        start = row.get("offset")

    if start is None:
        start = row.get("segment_start")

    start_value = float(start or 0.0)

    end = row.get("end")

    if end is None:
        end = row.get("segment_end")

    if end is None:
        duration = float(row.get("duration") or 0.0)
        end = start_value + duration

    return start_value, float(end)


def extract_text_from_reconciliation(
    reconciliation_path: Path,
) -> tuple[str, int]:
    document = json.loads(
        reconciliation_path.read_text(encoding="utf-8")
    )

    items = document.get("items", [])

    if not isinstance(items, list):
        raise ValueError(
            f"Invalid reconciliation items: {reconciliation_path}"
        )

    tokens: list[str] = []

    for item in items:
        token = (
            item.get("corrected_text")
            or item.get("selected_text")
            or item.get("text")
            or item.get("word")
            or item.get("en_text")
            or item.get("ar_text")
            or ""
        )

        token = str(token).strip()

        if token:
            tokens.append(token)

    text = " ".join(tokens)

    # Normalize spacing before punctuation while preserving Arabic text.
    text = re.sub(r"\s+([,.;:!?؟،؛])", r"\1", text)
    text = re.sub(r"([(\[{])\s+", r"\1", text)
    text = re.sub(r"\s+([)\]}])", r"\1", text)
    text = re.sub(r"\s+", " ", text).strip()

    return text, len(items)


def add_bottom_border(paragraph) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))

    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), "B7B7B7")
    borders.append(bottom)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    if "Review Transcript" not in styles:
        transcript_style = styles.add_style(
            "Review Transcript",
            WD_STYLE_TYPE.PARAGRAPH,
        )
    else:
        transcript_style = styles["Review Transcript"]

    transcript_style.font.name = "Arial"
    transcript_style.font.size = Pt(12)
    transcript_style.paragraph_format.space_after = Pt(8)
    transcript_style.paragraph_format.line_spacing = 1.15

    if "Segment Metadata" not in styles:
        metadata_style = styles.add_style(
            "Segment Metadata",
            WD_STYLE_TYPE.PARAGRAPH,
        )
    else:
        metadata_style = styles["Segment Metadata"]

    metadata_style.font.name = "Arial"
    metadata_style.font.size = Pt(9)
    metadata_style.font.color.rgb = None
    metadata_style.paragraph_format.space_after = Pt(3)


def build_document(
    *,
    source_audio_id: str,
    lecture_title: str,
    package_rows: list[dict[str, Any]],
    output_path: Path,
) -> None:
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.add_run(lecture_title)
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        "Gold Transcript Review — edit transcript paragraphs only"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    instructions = document.add_paragraph()
    instructions.add_run("Editing instructions: ").bold = True
    instructions.add_run(
        "Correct spelling, wording, punctuation, Arabic text, names, "
        "transliterations, Qur'anic quotations, and obvious ASR errors. "
        "Do not delete, rename, reorder, or edit segment metadata headings."
    )

    document.add_paragraph(
        f"Source audio ID: {source_audio_id}",
        style="Segment Metadata",
    )

    document.add_paragraph(
        f"Reviewable segments: {len(package_rows)}",
        style="Segment Metadata",
    )

    document.add_page_break()

    for index, row in enumerate(package_rows, start=1):
        segment_id = row["segment_id"]
        start_text = format_timestamp(row["segment_start"])
        end_text = format_timestamp(row["segment_end"])

        heading = document.add_paragraph()

        heading_run = heading.add_run(
            f"SEGMENT {index:04d} | {segment_id}"
        )
        heading_run.bold = True
        heading_run.font.size = Pt(11)

        metadata = document.add_paragraph(
            (
                f"Time: {start_text} → {end_text} | "
                f"Duration: {row['duration']:.3f}s | "
                f"Silver rows: {row['silver_row_count']}"
            ),
            style="Segment Metadata",
        )

        transcript = document.add_paragraph(
            row["silver_text"],
            style="Review Transcript",
        )

        transcript.paragraph_format.keep_together = True
        add_bottom_border(transcript)

        if index % 20 == 0 and index != len(package_rows):
            document.add_page_break()

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        f"{source_audio_id} — {SCHEMA_VERSION}"
    ).font.size = Pt(8)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()

    parser.add_argument(
        "--manifest",
        required=True,
    )
    parser.add_argument(
        "--reconciliation-dir",
        required=True,
    )
    parser.add_argument(
        "--output-dir",
        required=True,
    )
    parser.add_argument(
        "--source-audio-id",
        required=True,
    )
    parser.add_argument(
        "--lecture-title",
        required=True,
    )

    args = parser.parse_args()

    manifest_path = Path(args.manifest)
    reconciliation_dir = Path(args.reconciliation_dir)
    output_dir = Path(args.output_dir)

    manifest_rows = read_jsonl(manifest_path)

    package_rows: list[dict[str, Any]] = []
    missing_reconciliation: list[str] = []

    for sequence, row in enumerate(manifest_rows, start=1):
        segment_id = row["segment_id"]

        reconciliation_path = (
            reconciliation_dir
            / f"{safe_name(segment_id)}_reconciliation.json"
        )

        if not reconciliation_path.exists():
            missing_reconciliation.append(segment_id)
            continue

        silver_text, silver_row_count = (
            extract_text_from_reconciliation(
                reconciliation_path
            )
        )

        start, end = resolve_segment_bounds(row)
        duration = float(
            row.get("duration")
            or max(0.0, end - start)
        )

        package_rows.append({
            "review_sequence": sequence,
            "segment_id": segment_id,
            "audio_filepath": row.get("audio_filepath"),
            "source_audio_id": args.source_audio_id,
            "segment_start": start,
            "segment_end": end,
            "duration": duration,
            "silver_text": silver_text,
            "silver_row_count": silver_row_count,
            "reconciliation_path": str(
                reconciliation_path
            ),
        })

    if missing_reconciliation:
        raise RuntimeError(
            "Missing reconciliation files for:\n"
            + "\n".join(missing_reconciliation)
        )

    if not package_rows:
        raise RuntimeError(
            "No reviewable segment rows were produced."
        )

    output_dir.mkdir(parents=True, exist_ok=True)

    stem = f"{args.source_audio_id}_gold_review_v1"

    docx_path = output_dir / f"{stem}.docx"
    txt_path = output_dir / f"{stem}.txt"
    segments_path = output_dir / f"{stem}_segments.json"
    metadata_path = output_dir / f"{stem}_metadata.json"
    readme_path = output_dir / "README.md"

    build_document(
        source_audio_id=args.source_audio_id,
        lecture_title=args.lecture_title,
        package_rows=package_rows,
        output_path=docx_path,
    )

    with txt_path.open("w", encoding="utf-8") as handle:
        for row in package_rows:
            handle.write(
                f"[[SEGMENT_ID:{row['segment_id']}]]\n"
            )
            handle.write(
                (
                    f"[[TIME:"
                    f"{format_timestamp(row['segment_start'])}"
                    f"-"
                    f"{format_timestamp(row['segment_end'])}"
                    f"]]\n"
                )
            )
            handle.write(row["silver_text"] + "\n\n")

    segments_path.write_text(
        json.dumps(
            {
                "schema_version": SCHEMA_VERSION,
                "source_audio_id": args.source_audio_id,
                "segment_count": len(package_rows),
                "segments": package_rows,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )

    metadata = {
        "schema_version": SCHEMA_VERSION,
        "created_at": datetime.now(
            timezone.utc
        ).isoformat(),
        "source_audio_id": args.source_audio_id,
        "lecture_title": args.lecture_title,
        "manifest": str(manifest_path),
        "reconciliation_dir": str(reconciliation_dir),
        "segment_count": len(package_rows),
        "docx": str(docx_path),
        "txt": str(txt_path),
        "segments_sidecar": str(segments_path),
        "editing_rules": [
            "Edit transcript paragraphs only.",
            "Do not rename, delete, reorder, or modify segment headings.",
            "Do not combine two segment transcript paragraphs.",
            "Do not split one segment transcript into multiple segments.",
            "Track Changes may be enabled or disabled.",
            "Comments may be added, but corrections must appear in the transcript body.",
        ],
    }

    metadata_path.write_text(
        json.dumps(
            metadata,
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )

    readme_path.write_text(
        f"""# Gold Review Package

Source audio ID: `{args.source_audio_id}`

## File to edit

`{docx_path.name}`

## Editing rules

1. Edit only transcript paragraphs beneath each segment heading.
2. Do not delete or modify segment headings.
3. Do not reorder segments.
4. Do not merge adjacent segments.
5. Correct spelling, wording, punctuation, Arabic text, names, and transliterations.
6. Save the edited file under a new name, such as:

   `{args.source_audio_id}_gold_review_v1_EDITED.docx`

## Sidecar files

- `{segments_path.name}` preserves original Silver text and segment metadata.
- `{metadata_path.name}` describes the review package.
- `{txt_path.name}` is a plain-text reference export.

The edited DOCX will later be aligned against the Silver sidecar to produce Gold JSON, Gold JSONL, a delta report, and a NeMo training manifest.
""",
        encoding="utf-8",
    )

    print(
        json.dumps(
            {
                "docx": str(docx_path),
                "txt": str(txt_path),
                "segments": str(segments_path),
                "metadata": str(metadata_path),
                "readme": str(readme_path),
                "segment_count": len(package_rows),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
