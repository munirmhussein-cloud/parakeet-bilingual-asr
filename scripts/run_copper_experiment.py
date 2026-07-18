#!/usr/bin/env python3
"""Run one reproducible Copper decoding profile and export canonical 20s outputs."""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import platform
import subprocess
import sys
import time
from dataclasses import asdict, is_dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

WINDOW_SECONDS = 20.0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--config", required=True, type=Path)
    parser.add_argument("--lecture-id", required=True)
    parser.add_argument("--audio", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--download-root", type=Path, default=None)
    parser.add_argument("--overwrite", action="store_true")
    return parser.parse_args()


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def clean_text(value: Any) -> str:
    return " ".join(str(value or "").replace("\u00a0", " ").split())


def plain(value: Any) -> dict[str, Any]:
    if is_dataclass(value):
        return asdict(value)
    if hasattr(value, "_asdict"):
        return dict(value._asdict())
    if hasattr(value, "__dict__"):
        return dict(vars(value))
    return {}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(8 * 1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def gpu_name() -> str | None:
    result = subprocess.run(
        ["nvidia-smi", "--query-gpu=name", "--format=csv,noheader"],
        text=True,
        capture_output=True,
        check=False,
    )
    return result.stdout.strip().splitlines()[0] if result.returncode == 0 else None


def serialize_segments(iterator: Any) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for segment in iterator:
        raw = plain(segment)
        words = raw.get("words") or getattr(segment, "words", None) or []
        rows.append({
            "id": int(raw.get("id", getattr(segment, "id", len(rows)))),
            "seek": int(raw.get("seek", getattr(segment, "seek", 0))),
            "start": float(raw.get("start", getattr(segment, "start", 0.0))),
            "end": float(raw.get("end", getattr(segment, "end", 0.0))),
            "text": clean_text(raw.get("text", getattr(segment, "text", ""))),
            "tokens": list(raw.get("tokens", getattr(segment, "tokens", [])) or []),
            "temperature": raw.get("temperature", getattr(segment, "temperature", None)),
            "avg_logprob": raw.get("avg_logprob", getattr(segment, "avg_logprob", None)),
            "compression_ratio": raw.get("compression_ratio", getattr(segment, "compression_ratio", None)),
            "no_speech_prob": raw.get("no_speech_prob", getattr(segment, "no_speech_prob", None)),
            "words": [
                {
                    "start": float(plain(word).get("start", getattr(word, "start", 0.0))),
                    "end": float(plain(word).get("end", getattr(word, "end", 0.0))),
                    "word": str(plain(word).get("word", getattr(word, "word", ""))),
                    "probability": plain(word).get("probability", getattr(word, "probability", None)),
                }
                for word in words
            ],
        })
    return rows


def source_units(segments: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], str]:
    units: list[dict[str, Any]] = []
    for segment_index, segment in enumerate(segments):
        for word in segment.get("words") or []:
            if not str(word.get("word") or "").strip():
                continue
            units.append({
                "ordinal": len(units),
                "source_type": "word",
                "source_segment_index": segment_index,
                "start": float(word.get("start", segment["start"])),
                "end": float(word.get("end", segment["end"])),
                "text": str(word["word"]),
                "probability": word.get("probability"),
            })
    if units:
        return units, "word"

    for segment_index, segment in enumerate(segments):
        text = clean_text(segment.get("text"))
        if text:
            units.append({
                "ordinal": len(units),
                "source_type": "whisper_segment",
                "source_segment_index": segment_index,
                "start": float(segment["start"]),
                "end": float(segment["end"]),
                "text": text,
                "probability": None,
            })
    if not units:
        raise RuntimeError("No timestamped transcription units were produced")
    return units, "whisper_segment"


def canonical_rows(
    profile: str,
    lecture_id: str,
    audio: Path,
    duration: float,
    units: list[dict[str, Any]],
    source_type: str,
) -> list[dict[str, Any]]:
    count = max(1, math.ceil(duration / WINDOW_SECONDS))
    buckets: list[list[dict[str, Any]]] = [[] for _ in range(count)]
    for unit in units:
        midpoint = (float(unit["start"]) + float(unit["end"])) / 2.0
        index = max(0, min(count - 1, int(midpoint // WINDOW_SECONDS)))
        buckets[index].append(unit)

    assigned = [unit["ordinal"] for bucket in buckets for unit in bucket]
    if assigned != list(range(len(units))):
        raise RuntimeError("Canonical projection did not assign all units exactly once")

    rows: list[dict[str, Any]] = []
    for index, bucket in enumerate(buckets):
        start = index * WINDOW_SECONDS
        end = min(duration, (index + 1) * WINDOW_SECONDS)
        text = (
            " ".join("".join(str(unit["text"]) for unit in bucket).split())
            if source_type == "word"
            else " ".join(clean_text(unit["text"]) for unit in bucket)
        )
        rows.append({
            "schema_version": f"{profile}_canonical_20s_v1",
            "pipeline_version": profile,
            "lecture_id": lecture_id,
            "audio_filepath": str(audio),
            "segment_id": f"{lecture_id}__canonical_20s__{index:05d}",
            "segment_index": index,
            "segment_position": index,
            "segment_start": start,
            "segment_end": end,
            "duration": end - start,
            "has_copper_text": bool(text),
            "copper_text": text,
            "token_count": len(text.split()),
            "source_unit_type": source_type,
            "source_unit_count": len(bucket),
            "source_units": bucket,
        })
    return rows


def timestamp(seconds: float) -> str:
    total = max(0, int(seconds))
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def write_docx(path: Path, profile: str, lecture_id: str, rows: list[dict[str, Any]], package: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10.5)

    title = document.add_heading(lecture_id.replace("_", " ").title(), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{profile.upper()} — Faster-Whisper Turbo — Canonical 20s").bold = True
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(
        f"{len(rows):,} rows | {package['word_count']:,} words | "
        f"projection: {package['projection_source']}"
    ).italic = True

    for row in rows:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"[{timestamp(row['segment_start'])} – {timestamp(row['segment_end'])}] "
        )
        run.bold = True
        run.font.size = Pt(9)
        paragraph.add_run(row["copper_text"] or "[No speech]")
    document.save(path)


def main() -> int:
    args = parse_args()
    config = json.loads(args.config.expanduser().resolve().read_text(encoding="utf-8"))
    profile = str(config["profile"])
    audio = args.audio.expanduser().resolve()
    output_dir = args.output_dir.expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    if not audio.is_file():
        raise FileNotFoundError(audio)

    whole_json = output_dir / f"{args.lecture_id}_{profile}_faster_whisper_turbo_whole.json"
    canonical_json = output_dir / f"{args.lecture_id}_{profile}_canonical_20s.json"
    canonical_jsonl = output_dir / f"{args.lecture_id}_{profile}_canonical_20s.jsonl"
    docx_path = output_dir / f"{args.lecture_id}_{profile}_canonical_20s.docx"
    manifest_path = output_dir / "PACKAGE_MANIFEST.json"
    if not args.overwrite and any(path.exists() for path in (whole_json, canonical_json, canonical_jsonl, docx_path, manifest_path)):
        raise FileExistsError("Profile outputs exist; pass --overwrite")

    import ctranslate2
    import faster_whisper
    from faster_whisper import WhisperModel

    model_kwargs = {
        "device": "cuda",
        "compute_type": config.get("compute_type", "float16"),
        "num_workers": int(config.get("num_workers", 1)),
    }
    if args.download_root:
        args.download_root.mkdir(parents=True, exist_ok=True)
        model_kwargs["download_root"] = str(args.download_root.resolve())

    transcribe_options = dict(config["transcribe"])
    started = time.perf_counter()
    model = WhisperModel(config.get("model", "turbo"), **model_kwargs)
    iterator, info = model.transcribe(str(audio), **transcribe_options)
    segments = serialize_segments(iterator)
    elapsed = time.perf_counter() - started
    info_raw = plain(info)
    duration = float(info_raw.get("duration", getattr(info, "duration", 0.0)) or 0.0)
    transcript = " ".join(row["text"] for row in segments if row["text"])

    whole = {
        "schema_version": f"{profile}_faster_whisper_whole_v1",
        "pipeline_version": profile,
        "lecture_id": args.lecture_id,
        "created_at": utc_now(),
        "profile": config,
        "source_audio": {
            "path": str(audio),
            "size_bytes": audio.stat().st_size,
            "sha256": sha256_file(audio),
        },
        "engine": {
            "name": "faster-whisper",
            "model": config.get("model", "turbo"),
            "faster_whisper_version": faster_whisper.__version__,
            "ctranslate2_version": ctranslate2.__version__,
        },
        "runtime": {
            "elapsed_seconds": elapsed,
            "real_time_factor": elapsed / duration if duration else None,
            "python": sys.version,
            "platform": platform.platform(),
            "gpu": gpu_name(),
        },
        "language": info_raw.get("language", getattr(info, "language", None)),
        "language_probability": info_raw.get("language_probability", getattr(info, "language_probability", None)),
        "duration": duration,
        "segment_count": len(segments),
        "word_count": len(transcript.split()),
        "text": transcript,
        "segments": segments,
    }
    whole_json.write_text(json.dumps(whole, ensure_ascii=False, indent=2), encoding="utf-8")

    units, projection_source = source_units(segments)
    rows = canonical_rows(profile, args.lecture_id, audio, duration, units, projection_source)
    canonical_text = " ".join(row["copper_text"] for row in rows if row["copper_text"])
    package = {
        "schema_version": f"{profile}_canonical_20s_v1",
        "pipeline_version": profile,
        "lecture_id": args.lecture_id,
        "created_at": utc_now(),
        "profile": config,
        "source_whole_json": whole_json.name,
        "source_audio": whole["source_audio"],
        "engine": whole["engine"],
        "runtime": whole["runtime"],
        "duration": duration,
        "window_seconds": WINDOW_SECONDS,
        "projection_source": projection_source,
        "source_unit_count": len(units),
        "segment_count": len(rows),
        "nonempty_segment_count": sum(bool(row["copper_text"]) for row in rows),
        "word_count": len(canonical_text.split()),
        "text": canonical_text,
        "segments": rows,
    }
    canonical_json.write_text(json.dumps(package, ensure_ascii=False, indent=2), encoding="utf-8")
    with canonical_jsonl.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")
    write_docx(docx_path, profile, args.lecture_id, rows, package)

    manifest = {
        "schema_version": f"{profile}_package_manifest_v1",
        "pipeline_version": profile,
        "lecture_id": args.lecture_id,
        "profile": config,
        "files": {
            "whole_json": whole_json.name,
            "canonical_json": canonical_json.name,
            "canonical_jsonl": canonical_jsonl.name,
            "docx": docx_path.name,
        },
        "metrics": {
            "duration": duration,
            "elapsed_seconds": elapsed,
            "real_time_factor": whole["runtime"]["real_time_factor"],
            "canonical_segment_count": len(rows),
            "word_count": package["word_count"],
        },
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"\n=== {profile.upper()} COMPLETE ===")
    print(f"Output directory: {output_dir}")
    print(f"Detected language: {whole['language']}")
    print(f"Segments: {len(segments):,}")
    print(f"Canonical rows: {len(rows):,}")
    print(f"Words: {package['word_count']:,}")
    print(f"Elapsed seconds: {elapsed:.1f}")
    print(f"RTF: {whole['runtime']['real_time_factor']:.4f}")
    print(f"DOCX: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
