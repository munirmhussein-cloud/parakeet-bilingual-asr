
from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
import unicodedata
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


VIEW_PRIORITY = {
    "canonical_20s": 5,
    "context_10s_stride_5s": 4,
    "whole_slice": 3,
    "azure_pyannote_forced_ar": 2,
    "local_2p5s_contiguous": 1,
}

TOKEN_RE = re.compile(
    r"[\u0600-\u06ff]+|"
    r"[A-Za-z0-9]+(?:['’\-][A-Za-z0-9]+)*|"
    r"[^\w\s]",
    re.UNICODE,
)

PUNCTUATION = {
    ".", ",", "!", "?", ":", ";",
    "(", ")", "[", "]", "{", "}",
}


@dataclass(frozen=True)
class Observation:
    observation_id: str
    view: str
    source_id: str
    source_token_index: int
    surface: str
    normalized: str
    center: float
    start: float
    end: float
    timing_method: str


@dataclass
class Slot:
    slot_id: int
    observations: list[Observation] = field(default_factory=list)

    @property
    def center(self) -> float:
        if not self.observations:
            return 0.0

        values = sorted(
            observation.center
            for observation in self.observations
        )

        middle = len(values) // 2

        if len(values) % 2:
            return values[middle]

        return (
            values[middle - 1] + values[middle]
        ) / 2.0


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    rows = []

    with path.open("r", encoding="utf-8") as handle:
        for line_number, line in enumerate(handle, start=1):
            if not line.strip():
                continue

            try:
                row = json.loads(line)
            except json.JSONDecodeError as error:
                raise ValueError(
                    f"Invalid JSONL at {path}:{line_number}"
                ) from error

            if not isinstance(row, dict):
                raise TypeError(
                    f"Expected object at {path}:{line_number}"
                )

            rows.append(row)

    return rows


def write_jsonl(
    path: Path,
    rows: Iterable[dict[str, Any]],
) -> None:
    with path.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(
                json.dumps(
                    row,
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                )
                + "\n"
            )


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()

    with path.open("rb") as handle:
        for chunk in iter(
            lambda: handle.read(1024 * 1024),
            b"",
        ):
            digest.update(chunk)

    return digest.hexdigest()


def remove_arabic_diacritics(text: str) -> str:
    return "".join(
        character
        for character in unicodedata.normalize("NFD", text)
        if not (
            "\u064b" <= character <= "\u065f"
            or character == "\u0670"
            or "\u06d6" <= character <= "\u06ed"
        )
    )


def normalize_token(token: str) -> str:
    token = unicodedata.normalize("NFKC", token)
    token = remove_arabic_diacritics(token)
    token = token.casefold()

    token = (
        token
        .replace("ٱ", "ا")
        .replace("ى", "ي")
        .replace("ؤ", "و")
        .replace("ئ", "ي")
    )

    return re.sub(
        r"[^\w\u0600-\u06ff]+",
        "",
        token,
        flags=re.UNICODE,
    )


def tokenize(text: str) -> list[str]:
    return [
        match.group(0)
        for match in TOKEN_RE.finditer(text or "")
        if match.group(0).strip()
    ]


def is_punctuation(text: str) -> bool:
    return text in PUNCTUATION


def join_tokens(tokens: list[str]) -> str:
    output = ""

    no_leading_space = {
        ".", ",", "!", "?", ":", ";",
        "%", ")", "]", "}",
    }

    no_trailing_space = {
        "(", "[", "{",
    }

    for token in tokens:
        if not output:
            output = token
        elif token in no_leading_space:
            output += token
        elif output[-1:] in no_trailing_space:
            output += token
        else:
            output += " " + token

    return re.sub(r"\s+", " ", output).strip()


def intervals_overlap(
    left_start: float,
    left_end: float,
    right_start: float,
    right_end: float,
) -> bool:
    return (
        left_start < right_end
        and right_start < left_end
    )


def lexical_compatible(
    first: str,
    second: str,
) -> bool:
    if first == second:
        return True

    if not first or not second:
        return False

    if len(first) >= 4 and len(second) >= 4:
        if first.startswith(second) or second.startswith(first):
            return True

    return False


def safe_float(value: Any) -> float | None:
    if value is None or isinstance(value, bool):
        return None

    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def normalized_document_observations(
    *,
    document: dict[str, Any],
    view: str,
    segment_start: float,
    segment_end: float,
    clip_to_segment: bool,
) -> list[Observation]:
    source_id = str(
        document.get(
            "segment_id",
            document.get(
                "audio_id",
                f"{view}_source",
            ),
        )
    )

    window_start = float(
        document.get(
            "global_start",
            segment_start,
        )
    )

    window_end = float(
        document.get(
            "global_end",
            segment_end,
        )
    )

    words = document.get("words", [])

    if not isinstance(words, list):
        words = []

    surfaces = []

    for word in words:
        if not isinstance(word, dict):
            continue

        surface = str(
            word.get("text", "")
        ).strip()

        if surface:
            surfaces.append(
                (surface, word)
            )

    if not surfaces:
        surfaces = [
            (
                surface,
                {},
            )
            for surface in tokenize(
                str(document.get("text", ""))
            )
        ]

    count = len(surfaces)

    if count == 0:
        return []

    duration = max(
        window_end - window_start,
        0.001,
    )

    observations = []

    for index, (surface, word) in enumerate(surfaces):
        normalized = normalize_token(surface)

        if not normalized and not is_punctuation(surface):
            continue

        raw_start = safe_float(
            word.get("global_start")
        )

        raw_end = safe_float(
            word.get("global_end")
        )

        raw_center = None

        if raw_start is not None and raw_end is not None:
            raw_center = (
                raw_start + raw_end
            ) / 2.0
        elif raw_start is not None:
            raw_center = raw_start
        elif raw_end is not None:
            raw_center = raw_end

        raw_timing_credible = (
            raw_center is not None
            and window_start - 0.25
            <= raw_center
            <= window_end + 0.25
        )

        if raw_timing_credible:
            center = min(
                max(raw_center, window_start),
                window_end,
            )

            start = (
                raw_start
                if raw_start is not None
                else center
            )

            end = (
                raw_end
                if raw_end is not None
                else center
            )

            timing_method = "raw_global_timestamp"

        else:
            center = (
                window_start
                + ((index + 0.5) / count)
                * duration
            )

            token_width = duration / max(
                count,
                1,
            )

            start = max(
                window_start,
                center - token_width / 2.0,
            )

            end = min(
                window_end,
                center + token_width / 2.0,
            )

            timing_method = (
                "monotonic_window_interpolation"
            )

        if clip_to_segment:
            if not (
                segment_start
                <= center
                < segment_end
            ):
                continue

        observation_id = (
            f"{view}|{source_id}|{index}"
        )

        observations.append(
            Observation(
                observation_id=observation_id,
                view=view,
                source_id=source_id,
                source_token_index=index,
                surface=surface,
                normalized=normalized,
                center=round(center, 6),
                start=round(start, 6),
                end=round(end, 6),
                timing_method=timing_method,
            )
        )

    return observations


def azure_observations(
    *,
    row: dict[str, Any],
    segment_start: float,
    segment_end: float,
) -> list[Observation]:
    text = str(
        row.get(
            "azure_forced_ar_text",
            row.get(
                "text",
                "",
            ),
        )
    ).strip()

    tokens = tokenize(text)

    if not tokens:
        return []

    duration = max(
        segment_end - segment_start,
        0.001,
    )

    output = []

    for index, surface in enumerate(tokens):
        center = (
            segment_start
            + ((index + 0.5) / len(tokens))
            * duration
        )

        width = duration / len(tokens)

        output.append(
            Observation(
                observation_id=(
                    "azure_pyannote_forced_ar"
                    f"|{row.get('segment_id')}|{index}"
                ),
                view="azure_pyannote_forced_ar",
                source_id=str(
                    row.get("segment_id", "")
                ),
                source_token_index=index,
                surface=surface,
                normalized=normalize_token(surface),
                center=round(center, 6),
                start=round(
                    max(
                        segment_start,
                        center - width / 2.0,
                    ),
                    6,
                ),
                end=round(
                    min(
                        segment_end,
                        center + width / 2.0,
                    ),
                    6,
                ),
                timing_method=(
                    "segment_level_interpolation"
                ),
            )
        )

    return output


def assign_observations_to_slots(
    observations: list[Observation],
    *,
    lexical_window_seconds: float = 0.72,
    fallback_window_seconds: float = 0.24,
) -> list[Slot]:
    slots: list[Slot] = []

    for observation in sorted(
        observations,
        key=lambda item: (
            item.center,
            -VIEW_PRIORITY.get(
                item.view,
                0,
            ),
            item.source_id,
            item.source_token_index,
        ),
    ):
        best_slot = None
        best_score = None

        for slot in slots:
            distance = abs(
                observation.center
                - slot.center
            )

            same_lexeme = any(
                lexical_compatible(
                    observation.normalized,
                    existing.normalized,
                )
                for existing in slot.observations
            )

            threshold = (
                lexical_window_seconds
                if same_lexeme
                else fallback_window_seconds
            )

            if distance > threshold:
                continue

            # Do not merge two different tokens from the same source
            # unless they are identical punctuation observations.
            same_source_conflict = any(
                existing.view == observation.view
                and existing.source_id
                == observation.source_id
                and existing.source_token_index
                != observation.source_token_index
                for existing in slot.observations
            )

            if same_source_conflict:
                continue

            score = (
                0 if same_lexeme else 1,
                distance,
                slot.slot_id,
            )

            if best_score is None or score < best_score:
                best_score = score
                best_slot = slot

        if best_slot is None:
            best_slot = Slot(
                slot_id=len(slots)
            )

            slots.append(best_slot)

        best_slot.observations.append(
            observation
        )

    slots.sort(
        key=lambda slot: (
            slot.center,
            slot.slot_id,
        )
    )

    for index, slot in enumerate(slots):
        slot.slot_id = index

    return slots


def resolve_slot(
    slot: Slot,
) -> dict[str, Any]:
    by_normalized: dict[
        str,
        list[Observation]
    ] = defaultdict(list)

    for observation in slot.observations:
        key = (
            observation.normalized
            or observation.surface
        )

        by_normalized[key].append(
            observation
        )

    ranked = []

    for normalized, observations in by_normalized.items():
        views = {
            observation.view
            for observation in observations
        }

        ranked.append(
            (
                -len(views),
                -len(observations),
                -max(
                    VIEW_PRIORITY.get(
                        observation.view,
                        0,
                    )
                    for observation
                    in observations
                ),
                normalized,
                observations,
            )
        )

    ranked.sort()

    winning_observations = ranked[0][-1]

    surface_counter = Counter(
        observation.surface
        for observation in winning_observations
    )

    winning_surface = sorted(
        surface_counter.items(),
        key=lambda item: (
            -item[1],
            -max(
                VIEW_PRIORITY.get(
                    observation.view,
                    0,
                )
                for observation
                in winning_observations
                if observation.surface
                == item[0]
            ),
            item[0],
        ),
    )[0][0]

    views = sorted(
        {
            observation.view
            for observation
            in winning_observations
        }
    )

    source_ids = sorted(
        {
            observation.source_id
            for observation
            in winning_observations
        }
    )

    independent_view_count = len(views)

    if independent_view_count >= 2:
        tier = "A_corroborated"
    elif (
        len(source_ids) >= 2
        and views
        == ["context_10s_stride_5s"]
    ):
        tier = "B_stride_self_corroborated"
    else:
        tier = "C_single_witness"

    alternates = []

    for _, _, _, normalized, observations in ranked[1:]:
        alternates.append(
            {
                "normalized": normalized,
                "surfaces": sorted(
                    {
                        observation.surface
                        for observation
                        in observations
                    }
                ),
                "views": sorted(
                    {
                        observation.view
                        for observation
                        in observations
                    }
                ),
                "observation_ids": sorted(
                    observation.observation_id
                    for observation
                    in observations
                ),
            }
        )

    return {
        "slot_id": slot.slot_id,
        "center": round(
            slot.center,
            6,
        ),
        "text": winning_surface,
        "normalized": normalize_token(
            winning_surface
        ),
        "acceptance_tier": tier,
        "single_witness": (
            tier == "C_single_witness"
        ),
        "witness_view_count": (
            independent_view_count
        ),
        "views": views,
        "observation_count": len(
            winning_observations
        ),
        "observation_ids": sorted(
            observation.observation_id
            for observation
            in winning_observations
        ),
        "observations": [
            {
                "observation_id": (
                    observation.observation_id
                ),
                "view": observation.view,
                "source_id": (
                    observation.source_id
                ),
                "source_token_index": (
                    observation.source_token_index
                ),
                "surface": observation.surface,
                "normalized": (
                    observation.normalized
                ),
                "center": observation.center,
                "start": observation.start,
                "end": observation.end,
                "timing_method": (
                    observation.timing_method
                ),
            }
            for observation
            in sorted(
                winning_observations,
                key=lambda item: (
                    item.view,
                    item.source_id,
                    item.source_token_index,
                ),
            )
        ],
        "alternates": alternates,
    }


def immediate_duplicate_ngram_count(
    tokens: list[str],
    n: int = 6,
) -> int:
    normalized = [
        normalize_token(token)
        for token in tokens
        if normalize_token(token)
    ]

    count = 0

    for start in range(
        0,
        len(normalized) - (2 * n) + 1,
    ):
        first = normalized[
            start : start + n
        ]

        second = normalized[
            start + n : start + (2 * n)
        ]

        if first == second:
            count += 1

    return count


def format_timestamp(seconds: float) -> str:
    milliseconds = round(
        seconds * 1000
    )

    hours, remainder = divmod(
        milliseconds,
        3_600_000,
    )

    minutes, remainder = divmod(
        remainder,
        60_000,
    )

    seconds, milliseconds = divmod(
        remainder,
        1000,
    )

    return (
        f"{hours:02d}:"
        f"{minutes:02d}:"
        f"{seconds:02d}."
        f"{milliseconds:03d}"
    )


def export_docx(
    rows: list[dict[str, Any]],
    path: Path,
    *,
    title_text: str,
) -> None:
    document = Document()

    document.styles[
        "Normal"
    ].font.name = "Arial"

    document.styles[
        "Normal"
    ].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    subtitle.add_run(
        "Align-then-vote reconciliation; "
        "canonical 20-second output"
    ).bold = True

    document.add_paragraph(
        "Tier A: corroborated by multiple views. "
        "Tier B: corroborated across overlapping stride windows. "
        "Tier C: retained single-witness evidence."
    )

    document.add_paragraph("")

    for row in rows:
        heading = document.add_paragraph()

        heading.add_run(
            f"SEGMENT "
            f"{row['segment_position'] + 1:04d}"
            f" | {row['segment_id']}"
        ).bold = True

        metadata = document.add_paragraph()

        metadata.add_run(
            "Time: "
            f"{format_timestamp(row['segment_start'])}"
            " → "
            f"{format_timestamp(row['segment_end'])}"
            " | Tokens: "
            f"{row['token_count']}"
            " | A/B/C: "
            f"{row['tier_counts'].get('A_corroborated', 0)}"
            "/"
            f"{row['tier_counts'].get('B_stride_self_corroborated', 0)}"
            "/"
            f"{row['tier_counts'].get('C_single_witness', 0)}"
        )

        document.add_paragraph(
            row["silver_text"]
        )

        document.add_paragraph("")

    document.save(path)


def main() -> int:
    parser = argparse.ArgumentParser()

    parser.add_argument(
        "--lecture-id",
        required=True,
    )

    parser.add_argument(
        "--whole",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--canonical",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--context",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--local",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--canonical-manifest",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--azure-parent",
        type=Path,
    )

    parser.add_argument(
        "--output-dir",
        type=Path,
        required=True,
    )

    parser.add_argument(
        "--output-prefix",
        required=True,
    )

    parser.add_argument(
        "--schema-version",
        required=True,
    )

    parser.add_argument(
        "--title",
        required=True,
    )

    args = parser.parse_args()

    args.output_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    whole_rows = read_jsonl(
        args.whole
    )

    canonical_rows = read_jsonl(
        args.canonical
    )

    context_rows = read_jsonl(
        args.context
    )

    local_rows = read_jsonl(
        args.local
    )

    manifest_rows = read_jsonl(
        args.canonical_manifest
    )

    azure_by_segment = {}

    if args.azure_parent:
        for row in read_jsonl(
            args.azure_parent
        ):
            azure_by_segment[
                str(row["segment_id"])
            ] = row

    if len(whole_rows) != 1:
        raise ValueError(
            "Expected one whole-view document."
        )

    if len(manifest_rows) < len(
        canonical_rows
    ):
        raise ValueError(
            "Canonical manifest is shorter "
            "than canonical normalized input."
        )

    whole_document = whole_rows[0]

    output_rows = []
    provenance_rows = []

    all_input_observation_ids = set()
    all_accounted_observation_ids = set()

    immediate_duplicate_6grams = 0
    chronology_errors = []

    previous_segment_end = None

    for position, canonical in enumerate(
        canonical_rows
    ):
        manifest = manifest_rows[
            position
        ]

        segment_id = str(
            manifest.get(
                "segment_id",
                canonical.get(
                    "segment_id",
                    f"segment_{position:06d}",
                ),
            )
        )

        segment_start = float(
            manifest.get(
                "segment_start",
                canonical.get(
                    "global_start",
                    position * 20.0,
                ),
            )
        )

        segment_end = float(
            manifest.get(
                "segment_end",
                canonical.get(
                    "global_end",
                    segment_start + 20.0,
                ),
            )
        )

        duration = float(
            manifest.get(
                "duration",
                segment_end - segment_start,
            )
        )

        if (
            previous_segment_end is not None
            and segment_start
            < previous_segment_end - 0.05
        ):
            chronology_errors.append(
                {
                    "segment_id": segment_id,
                    "segment_start": (
                        segment_start
                    ),
                    "previous_segment_end": (
                        previous_segment_end
                    ),
                }
            )

        previous_segment_end = segment_end

        observations = []

        observations.extend(
            normalized_document_observations(
                document=whole_document,
                view="whole_slice",
                segment_start=segment_start,
                segment_end=segment_end,
                clip_to_segment=True,
            )
        )

        observations.extend(
            normalized_document_observations(
                document=canonical,
                view="canonical_20s",
                segment_start=segment_start,
                segment_end=segment_end,
                clip_to_segment=True,
            )
        )

        for document in context_rows:
            if intervals_overlap(
                segment_start,
                segment_end,
                float(
                    document.get(
                        "global_start",
                        segment_start,
                    )
                ),
                float(
                    document.get(
                        "global_end",
                        segment_end,
                    )
                ),
            ):
                observations.extend(
                    normalized_document_observations(
                        document=document,
                        view=(
                            "context_10s_stride_5s"
                        ),
                        segment_start=segment_start,
                        segment_end=segment_end,
                        clip_to_segment=True,
                    )
                )

        for document in local_rows:
            if intervals_overlap(
                segment_start,
                segment_end,
                float(
                    document.get(
                        "global_start",
                        segment_start,
                    )
                ),
                float(
                    document.get(
                        "global_end",
                        segment_end,
                    )
                ),
            ):
                observations.extend(
                    normalized_document_observations(
                        document=document,
                        view=(
                            "local_2p5s_contiguous"
                        ),
                        segment_start=segment_start,
                        segment_end=segment_end,
                        clip_to_segment=True,
                    )
                )

        if segment_id in azure_by_segment:
            observations.extend(
                azure_observations(
                    row=azure_by_segment[
                        segment_id
                    ],
                    segment_start=segment_start,
                    segment_end=segment_end,
                )
            )

        for observation in observations:
            all_input_observation_ids.add(
                observation.observation_id
            )

        slots = assign_observations_to_slots(
            observations
        )

        resolved_tokens = [
            resolve_slot(slot)
            for slot in slots
        ]

        for token in resolved_tokens:
            for observation in token[
                "observations"
            ]:
                all_accounted_observation_ids.add(
                    observation[
                        "observation_id"
                    ]
                )

            for alternate in token[
                "alternates"
            ]:
                for observation_id in alternate[
                    "observation_ids"
                ]:
                    all_accounted_observation_ids.add(
                        observation_id
                    )

        surfaces = [
            token["text"]
            for token in resolved_tokens
        ]

        silver_text = join_tokens(
            surfaces
        )

        duplicate_count = (
            immediate_duplicate_ngram_count(
                surfaces,
                n=6,
            )
        )

        immediate_duplicate_6grams += (
            duplicate_count
        )

        tier_counts = dict(
            Counter(
                token["acceptance_tier"]
                for token in resolved_tokens
            )
        )

        view_contribution_counts = dict(
            Counter(
                view
                for token in resolved_tokens
                for view in token["views"]
            )
        )

        output_row = {
            "schema_version": (
                args.schema_version
            ),
            "lecture_id": (
                args.lecture_id
            ),
            "segment_position": (
                position
            ),
            "segment_index": position,
            "segment_id": segment_id,
            "audio_filepath": str(
                manifest.get(
                    "audio_filepath",
                    canonical.get(
                        "audio_filepath",
                        "",
                    ),
                )
            ),
            "segment_start": round(
                segment_start,
                6,
            ),
            "segment_end": round(
                segment_end,
                6,
            ),
            "duration": round(
                duration,
                6,
            ),
            "silver_text": silver_text,
            "has_silver_text": bool(
                silver_text
            ),
            "token_count": len(
                resolved_tokens
            ),
            "tier_counts": tier_counts,
            "view_contribution_counts": (
                view_contribution_counts
            ),
            "immediate_duplicate_6gram_count": (
                duplicate_count
            ),
            "tokens": resolved_tokens,
            "reconciliation": {
                "method": (
                    "time_keyed_multiview_"
                    "align_then_vote_v1"
                ),
                "single_witness_policy": (
                    "accept_and_flag"
                ),
                "concatenation_used": False,
                "segment_source_selection_used": (
                    False
                ),
                "human_review_used": False,
            },
        }

        output_rows.append(
            output_row
        )

        provenance_rows.extend(
            {
                "schema_version": (
                    args.schema_version
                    + "_token_provenance"
                ),
                "segment_position": position,
                "segment_id": segment_id,
                **token,
            }
            for token in resolved_tokens
        )

    unaccounted = sorted(
        all_input_observation_ids
        - all_accounted_observation_ids
    )

    output_jsonl = (
        args.output_dir
        / f"{args.output_prefix}_segment_level.jsonl"
    )

    provenance_jsonl = (
        args.output_dir
        / f"{args.output_prefix}_token_provenance.jsonl"
    )

    output_json = (
        args.output_dir
        / f"{args.output_prefix}.json"
    )

    report_json = (
        args.output_dir
        / f"{args.output_prefix}_report.json"
    )

    docx_path = (
        args.output_dir
        / f"{args.output_prefix}.docx"
    )

    write_jsonl(
        output_jsonl,
        output_rows,
    )

    write_jsonl(
        provenance_jsonl,
        provenance_rows,
    )

    output_json.write_text(
        json.dumps(
            {
                "schema_version": (
                    args.schema_version
                    + "_export"
                ),
                "lecture_id": (
                    args.lecture_id
                ),
                "segment_count": len(
                    output_rows
                ),
                "segments": output_rows,
            },
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        + "\n",
        encoding="utf-8",
    )

    export_docx(
        output_rows,
        docx_path,
        title_text=args.title,
    )

    validation = {
        "segment_positions_ordered": (
            [
                row["segment_position"]
                for row in output_rows
            ]
            == list(
                range(len(output_rows))
            )
        ),
        "segment_ids_unique": (
            len(
                {
                    row["segment_id"]
                    for row in output_rows
                }
            )
            == len(output_rows)
        ),
        "chronology_error_count": len(
            chronology_errors
        ),
        "zero_drop_invariant": (
            len(unaccounted) == 0
        ),
        "unaccounted_observation_count": (
            len(unaccounted)
        ),
        "immediate_duplicate_6gram_count": (
            immediate_duplicate_6grams
        ),
        "all_segments_populated": all(
            row["has_silver_text"]
            for row in output_rows
        ),
    }

    validation["passed"] = (
        validation[
            "segment_positions_ordered"
        ]
        and validation[
            "segment_ids_unique"
        ]
        and validation[
            "chronology_error_count"
        ]
        == 0
        and validation[
            "zero_drop_invariant"
        ]
        and validation[
            "immediate_duplicate_6gram_count"
        ]
        == 0
    )

    report = {
        "schema_version": (
            args.schema_version
            + "_report"
        ),
        "lecture_id": args.lecture_id,
        "segment_count": len(
            output_rows
        ),
        "total_tokens": sum(
            row["token_count"]
            for row in output_rows
        ),
        "tier_distribution": dict(
            Counter(
                token["acceptance_tier"]
                for row in output_rows
                for token in row["tokens"]
            )
        ),
        "single_witness_token_count": sum(
            token["single_witness"]
            for row in output_rows
            for token in row["tokens"]
        ),
        "outputs": {
            "segment_jsonl": str(
                output_jsonl
            ),
            "token_provenance_jsonl": str(
                provenance_jsonl
            ),
            "json": str(output_json),
            "docx": str(docx_path),
        },
        "sha256": {
            "segment_jsonl": sha256_file(
                output_jsonl
            ),
            "token_provenance_jsonl": (
                sha256_file(
                    provenance_jsonl
                )
            ),
            "json": sha256_file(
                output_json
            ),
            "docx": sha256_file(
                docx_path
            ),
        },
        "chronology_errors": (
            chronology_errors
        ),
        "unaccounted_observations": (
            unaccounted
        ),
        "validation": validation,
    }

    report_json.write_text(
        json.dumps(
            report,
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
        )
        + "\n",
        encoding="utf-8",
    )

    print(
        json.dumps(
            report,
            ensure_ascii=False,
            indent=2,
        )
    )

    if not validation["passed"]:
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
